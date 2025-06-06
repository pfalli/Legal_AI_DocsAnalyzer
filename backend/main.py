from fastapi import FastAPI, Depends, UploadFile, File
from . import models
from .database import engine, get_db
from sqlalchemy.orm import Session
import shutil
import os
import docx
import pdfplumber
from fastapi.middleware.cors import CORSMiddleware
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Create database tables
# This should ideally be handled by migrations (e.g., Alembic) in a production app
models.Base.metadata.create_all(bind=engine)

app = FastAPI()

origins = [
    # "http://localhost:3000", # Old Next.js frontend (can be removed or kept if needed)
    "http://localhost:5173", # Your new Vite frontend
    # You can add other origins if needed
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"], # Allows all methods
    allow_headers=["*"], # Allows all headers
)

# Define a directory to store uploaded files (optional, for temporary storage or inspection)
UPLOAD_DIRECTORY = "./uploaded_files"
if not os.path.exists(UPLOAD_DIRECTORY):
    os.makedirs(UPLOAD_DIRECTORY)

def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document as PythonDocxDocument # Local import to ensure it's clear
        doc = PythonDocxDocument(file_path) # Use file_path directly
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return "" 

def extract_text_from_pdf(file_path: str) -> str:
    full_text = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            full_text.append(page.extract_text() or "") # Ensure to handle None from extract_text
    return '\n'.join(full_text)

async def extract_clauses_with_openai(text: str) -> dict:
    """
    Uses OpenAI's GPT to extract predefined legal clauses from text.
    """
    if not client:
        return {"error": "OpenAI API key not configured."}

    clauses_to_extract = [
        "confidentiality",
        "termination",
        "governing_law",
        "dispute_resolution",
        "force_majeure",
        "indemnity",
        "payment_terms",
        "intellectual_property",
        "non_compete",
        "notice"
    ]

    prompt = f"""
    Analyze the legal document text provided below and extract the following information:

    1.  **document_type**: Classify the type of the document (e.g., NDA, Employment Contract, Lease Agreement, Privacy Policy, Service Agreement, Other). If "Other", specify.

    2.  **metadata**: Extract the following details. If a detail is not found, indicate "Not found" or null.
        -   title: The main title of the document.
        -   parties_involved: A list of all named parties (individuals or organizations).
        -   effective_date: The date the agreement becomes effective.
        -   expiration_date: The date the agreement expires, if specified.
        -   place_of_signing: The location where the document was signed, if specified.
        -   date_of_signing: The date the document was signed.
        -   governing_law: The jurisdiction whose laws govern the agreement (e.g., "State of California", "England and Wales").
        -   signatures: A list of names or roles of the signatories.

    3.  **clauses**: For each of the following clause types, extract the full text of the clause if present. If a clause is not found, indicate "Not found".
        Clause types to extract: {', '.join(clauses_to_extract)}.

    Document Text:
    \"\"\"
    {text[:4000]} 
    \"\"\"

    Respond in a single JSON object with three top-level keys: "document_type", "metadata", and "clauses".
    The "metadata" key should have a nested JSON object for its values.
    The "clauses" key should have a nested JSON object where keys are the clause names.

    Example of the expected JSON response format:
    {{
      "document_type": "Non-Disclosure Agreement",
      "metadata": {{
        "title": "Mutual Non-Disclosure Agreement",
        "parties_involved": ["Innovate Corp", "Alpha Solutions Inc."],
        "effective_date": "2025-06-05",
        "expiration_date": "2027-06-04",
        "place_of_signing": "Not found",
        "date_of_signing": "2025-06-01",
        "governing_law": "State of Delaware",
        "signatures": ["Jane Doe (CEO, Innovate Corp)", "John Smith (Director, Alpha Solutions Inc.)"]
      }},
      "clauses": {{
        "confidentiality": "Each party agrees to keep confidential all proprietary information...",
        "termination": "This Agreement may be terminated by either party upon 30 days written notice...",
        "governing_law": "This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware.",
        "dispute_resolution": "Not found"
        // ... other clauses ...
      }}
    }}
    """
    # Using text[:4000] as a simple way to limit token usage for now.
    # For longer documents, you'll need a more sophisticated chunking strategy.

    try:
        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are an AI assistant specialized in legal document analysis."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2, # Lower temperature for more deterministic output
        )
        # Ensure the response is in the expected format and parse it
        # This is a simplified parsing, robust error handling would be needed
        content = response.choices[0].message.content
        # Attempt to parse the string content as JSON
        import json
        try:
            extracted_clauses = json.loads(content)
        except json.JSONDecodeError:
            extracted_clauses = {"error": "Failed to parse OpenAI response as JSON", "raw_response": content}
        return extracted_clauses
    except Exception as e:
        return {"error": f"OpenAI API call failed: {str(e)}"}

@app.get("/")
def root() -> dict[str, str]:
    return {"message": "Hello ClauseMate Backend"}


# About page route
@app.get("/about")
def about() -> dict[str, str]:
    return {"message": "This is the about page for ClauseMate."}


# Updated Route to upload and process a document
@app.post("/documents/upload/")
async def upload_document(db: Session = Depends(get_db), file: UploadFile = File(...)):
    original_filename = file.filename
    file_extension = original_filename.split(".")[-1].lower()
    
    temp_file_path = os.path.join(UPLOAD_DIRECTORY, original_filename)
    with open(temp_file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    extracted_text = ""
    clauses = {} # Initialize clauses dictionary
    try:
        if file_extension == "docx":
            extracted_text = extract_text_from_docx(temp_file_path)
        elif file_extension == "pdf":
            extracted_text = extract_text_from_pdf(temp_file_path)
        else:
            os.remove(temp_file_path)
            return {"error": "Unsupported file type. Please upload DOCX or PDF."}
        
        if extracted_text:
            # Call OpenAI to extract clauses
            clauses = await extract_clauses_with_openai(extracted_text)
        else:
            clauses = {"message": "No text could be extracted from the document to find clauses."}

        db_document = models.Document(
            filename=original_filename, 
            content_text=extracted_text, # Storing full text
            document_type="Pending Classification" 
        )
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        os.remove(temp_file_path)
        
        return {
            "id": db_document.id,
            "filename": db_document.filename,
            "message": "File uploaded and processed successfully.",
            "extracted_text_preview": (extracted_text[:200] + "...") if extracted_text else "No text extracted.",
            "extracted_text": extracted_text,  # <-- Add this line
            "extracted_clauses": clauses # Add clauses to the response
        }
    except Exception as e:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        # Consider logging the full exception e for debugging
        return {"error": f"Failed to process file: {str(e)}", "details": str(e)}

# Example: Route to create a document (illustrative)
@app.post("/documents/")
def create_document_stub(filename: str, db: Session = Depends(get_db)):
    # This is a placeholder. We'll implement actual document creation later.
    # For now, let's just simulate creating a DB entry.
    db_document = models.Document(filename=filename, content_text="Sample content...")
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    return db_document


# Example: Route to list documents (illustrative)
@app.get("/documents/")
def list_documents_stub(db: Session = Depends(get_db)):
    documents = db.query(models.Document).all()
    return documents
